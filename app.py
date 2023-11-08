from flask import Flask, send_file, render_template, request, redirect, url_for, flash, jsonify
from flask_sqlalchemy import SQLAlchemy
from flask_migrate import Migrate  
from flask import send_file
import openpyxl
from openpyxl.utils import get_column_letter
from io import BytesIO
from docx import Document
import base64
from docx.shared import Inches
from sqlalchemy import asc
from collections import defaultdict
import pandas as pd
import os
import shutil
from collections import OrderedDict
import requests
import urllib3
from multiprocessing import Pool
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)


# Move the app creation into a factory function.
db = SQLAlchemy()
migrate = Migrate()

# Mapping current categories to new categories
category_mapping = {
    'Vehicle Driver': 'Private Vehicle',
    'Train': 'Public Transport',
    'Vehicle Passenger': 'Private Vehicle',
    'Motorcycle': 'Private Vehicle',
    'Public Bus': 'Public Transport',
    'Walking': 'Walking',
    'Bicycle': 'Bicycle',
    'Tram': 'Public Transport',
    'Taxi': 'Private Vehicle',
    'Other': 'Other Modes',
    'Jogging': 'Other Modes',
    'School Bus': 'Private Vehicle',
    'Mobility Scooter': 'Other Modes'
}

def process_single_vista_data(journey_type):
    try:
        df = fetch_vista_data(journey_type)
        if df.empty:
            print("No data fetched from API.")
            return {}
        
        mode_col_mapping = {
            "WORK": 'jtwmode',
            "EDUCATION": 'jtemode',
            "STOPS": 'mainmode',
            "TRIPS": 'linkmode'
        }

        mode_col = mode_col_mapping[journey_type]

        dist_col_mapping = {
            "WORK": 'jtwdist',
            "EDUCATION": 'jtedist',
            "STOPS": 'vistadist',
            "TRIPS": 'cumdist'
        }

        distance_col = dist_col_mapping[journey_type]

        time_col_mapping = {
            "WORK": 'jtwelapsedtime',
            "EDUCATION": 'jteelapsedtime',
            "STOPS": 'travtime',
            "TRIPS": 'travtime'
        }

        time_col = time_col_mapping[journey_type]

        df[time_col] = pd.to_numeric(df[time_col], errors='coerce')
        df[distance_col] = pd.to_numeric(df[distance_col], errors='coerce')

        df.dropna(subset=[time_col, distance_col, mode_col], inplace=True)
        
        # Convert elapsed time to hours if data is in minutes
        df[time_col] = df[time_col] / 60.0
        
        # Assuming a default trip value of 1 for each row
        df['trips'] = 1

        df[mode_col] = df[mode_col].map(category_mapping)

        # Aggregate data
        mode_raw_trips = df.groupby(mode_col)['trips'].sum().to_dict()
        mode_raw_distance = df.groupby(mode_col)[distance_col].sum().to_dict()
        mode_raw_time = df.groupby(mode_col)[time_col].sum().to_dict()

        # Collect summary data
        total_number_of_trips = df['trips'].sum()
        total_distance = df[distance_col].sum()
        total_time = df[time_col].sum()

        summary = {
            'mode_raw_trips': mode_raw_trips,
            'total_number_of_trips': total_number_of_trips,
            'mode_raw_distance': mode_raw_distance,
            'mode_raw_time': mode_raw_time,
            'unique_modes': df[mode_col].unique().tolist(),
            'total_distance': total_distance,
            'total_time': total_time
        }

        return summary
    except Exception as e:
        print("Error inside process_vista_data:", e)
        return {}

def process_vista_data():
    journey_types = ["TRIPS", "STOPS", "WORK", "EDUCATION"]

    # Set a consistent order based on the modes from the category_mapping
    mode_order = list(dict.fromkeys(category_mapping.values()))

    # Initialize the combined summary
    combined_summary = {
        'total_number_of_trips': 0,
        'total_distance': 0,
        'total_time': 0,
        'unique_modes': set()
    }

    # Initialize accumulators for raw counts
    mode_raw_trips = {mode: 0 for mode in mode_order}
    mode_raw_distance = {mode: 0 for mode in mode_order}
    mode_raw_time = {mode: 0 for mode in mode_order}

    with Pool() as pool:
        summaries = pool.map(process_single_vista_data, journey_types)

    # Accumulate raw counts from summaries
    for summary in summaries:
        combined_summary['total_number_of_trips'] += summary['total_number_of_trips']
        combined_summary['total_distance'] += summary['total_distance']
        combined_summary['total_time'] += summary['total_time']
        combined_summary['unique_modes'].update(summary['unique_modes'])
        
        for mode, count in summary['mode_raw_trips'].items():
            mode_raw_trips[mode] += count
        for mode, distance in summary['mode_raw_distance'].items():
            mode_raw_distance[mode] += distance
        for mode, time in summary['mode_raw_time'].items():
            mode_raw_time[mode] += time

    # Calculate percentages from raw counts
    combined_summary['mode_share_trips'] = {mode: (count / combined_summary['total_number_of_trips'] * 100) if combined_summary['total_number_of_trips'] else 0 for mode, count in mode_raw_trips.items()}
    combined_summary['mode_share_distance'] = {mode: (distance / combined_summary['total_distance'] * 100) if combined_summary['total_distance'] else 0 for mode, distance in mode_raw_distance.items()}
    combined_summary['mode_share_time'] = {mode: (time / combined_summary['total_time'] * 100) if combined_summary['total_time'] else 0 for mode, time in mode_raw_time.items()}

    # Convert unique modes from set to a sorted list based on the predefined mode_order
    combined_summary['unique_modes'] = sorted(list(combined_summary['unique_modes']), key=lambda x: mode_order.index(x))

    print(combined_summary['unique_modes'])
    print(combined_summary)
    return combined_summary


def fetch_vista_data(journey_type):
    # Define URL and parameters
    url = 'https://discover.data.vic.gov.au/api/3/action/datastore_search'
    limit = 32000  # Maximum number of records per request
    offset = 0     # Starting point for the results

    resource_ids = {
        "WORK": '5731a38c-6931-4135-bb95-78a9c601a9ab',
        "EDUCATION": '6967fe21-969f-4a53-8997-60b3b9126067',
        "STOPS": '79fd0247-ef20-4145-a431-3577cafe5c41',
        "TRIPS": '32949433-4c83-4ce8-83f5-7f7b40bd04d0'
    }
    resource_id = resource_ids[journey_type]

    all_records = []

    while True:
        params = {
            'resource_id': resource_id,
            'limit': limit,
            'offset': offset
        }

        # Fetch data from the API
        response = requests.get(url, params=params, verify=False)

        # Check if the request was successful
        if response.status_code == 200:
            data = response.json()
            records = data['result']['records']

            if not records:  # Break the loop if no more records are found
                break

            all_records.extend(records)

            # Increase the offset by the number of records fetched to fetch the next page in the next iteration
            offset += len(records)
        else:
            # Handle the error
            print("Error fetching VISTA data:", response.status_code)
            break

    return pd.DataFrame(all_records)


def create_app():
    app = Flask(__name__)
    app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///climatedashboard.db'
    UPLOAD_FOLDER = 'uploads'
    app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
    ALLOWED_EXTENSIONS = {'csv'}

    print("Database URI:", app.config['SQLALCHEMY_DATABASE_URI'])  # Add this line for debugging
    app.secret_key = 'your_secret_key_here'

    db.init_app(app)
    migrate.init_app(app, db)

    from models import User, ReportingPeriod, EmissionReport, BusinessOrFacility, EnergyRatingData, EmissionFactor

    def get_original_filename(file_path):
        if os.path.exists(file_path):
            with open(file_path, 'r') as f:
                return f.read()
        return "No file chosen"

    @app.route('/')
    def index():
        return render_template('index.html')
    
    @app.route('/business_index')
    def business_index():
        return render_template('business_and_industry/business_index.html')
    
    @app.route('/esd_index')
    def esd_index():
        emission_factors = EmissionFactor.query.all()

        # Dictionary to store total emissions reduction per year
        total_emissions_reduction_per_year = {}

        for ef in emission_factors:
            year = ef.year
            # Get all the EnergyRatingData records for the current emission factor's year
            energy_data_for_year = EnergyRatingData.query.filter_by(emission_factor_id=ef.id).all()
            total_reduction = sum(data.emissions_reduction for data in energy_data_for_year if data.emissions_reduction is not None)
            total_emissions_reduction_per_year[year] = total_reduction

        return render_template('esd/esd_index.html', emission_factors=emission_factors, total_emissions_reduction=total_emissions_reduction_per_year)

    
    @app.route('/api/all_reporting_periods', methods=['GET'])
    def all_reporting_periods():
        periods = ReportingPeriod.query.all()
        return jsonify([{
            'id': period.id,
            'start_month': period.start_month,
            'start_year': period.start_year,
            'end_month': period.end_month,
            'end_year': period.end_year
        } for period in periods])
    
    @app.route('/api/emissions_over_time', methods=['GET'])
    def get_emissions_over_time():
        periods = ReportingPeriod.query.all()
        data = []
        for period in periods:
            solar_emissions = sum([report.co2_emissions_solar for report in period.reports])
            other_emissions = sum([report.other_emissions for report in period.reports])
            data.append({
                "end_month": period.end_month, 
                "end_year": period.end_year,   
                "solar_emissions": solar_emissions,
                "other_emissions": other_emissions
            })
        return jsonify(data)

    @app.route('/create_reporting_period', methods=['GET', 'POST'])
    def create_reporting_period():
        # this creates a new reporting period, where the user can enter the start and end month and year
        if request.method == 'POST':
            start_month = request.form['start_month']
            start_year = request.form['start_year']
            end_month = request.form['end_month']
            end_year = request.form['end_year']
            reporting_period = ReportingPeriod(start_month=start_month, start_year=start_year, end_month=end_month, end_year=end_year)
            db.session.add(reporting_period)
            db.session.commit()
            return redirect(url_for('business_index'))
        return render_template('business_and_industry/create_reporting_period.html')

    @app.route('/remove_reporting_period', methods=['GET', 'POST'])
    def remove_reporting_period():
        if request.method == 'POST':
            period_id = request.form.get('reporting_period_id')
            period = ReportingPeriod.query.get(period_id)
            
            if period:
                db.session.delete(period)
                db.session.commit()
                return redirect(url_for('business_index'))
            else:
                return "Error: Reporting Period Not Found", 400

        # If GET request, display the page with all reporting periods
        periods = ReportingPeriod.query.all()
        return render_template('business_and_industry/remove_reporting_period.html', periods=periods)
    
    def get_reporting_period_by_id(period_id):
        return ReportingPeriod.query.get(period_id)

    @app.route('/modify_reporting_period', methods=['GET', 'POST'])
    def modify_reporting_period():
        if request.method == 'POST':
            period_id = request.form['reporting_period']
            start_month = request.form['start_month']
            start_year = request.form['start_year']
            end_month = request.form['end_month']
            end_year = request.form['end_year']
            
            reporting_period = get_reporting_period_by_id(period_id)
            reporting_period.start_month = start_month
            reporting_period.start_year = start_year
            reporting_period.end_month = end_month
            reporting_period.end_year = end_year

            db.session.commit()
            return redirect(url_for('business_index'))
        all_periods = get_all_reporting_periods()
        return render_template('business_and_industry/modify_reporting_period.html', all_periods=all_periods)
    
    def get_all_reporting_periods():
        """Fetches all reporting periods from the database, sorted by start year and month."""
        return ReportingPeriod.query.order_by(ReportingPeriod.start_year, ReportingPeriod.start_month).all()

    @app.route('/view_reports', methods=['GET', 'POST'])
    def view_reports():
        reporting_periods = ReportingPeriod.query.all()

        if not reporting_periods:  # Check if there are no reporting periods
            return render_template('business_and_industry/view_reports.html', reporting_periods=[], reports=[])

        if request.method == 'POST':
            reporting_period_id = request.form['reporting_period']
        else:  # If method is GET, default to the first reporting period
            reporting_period_id = reporting_periods[0].id

        selected_reporting_period = ReportingPeriod.query.get(reporting_period_id)
        reports = selected_reporting_period.reports

        # Calculations
        total_ghg_savings = sum([report.total_emissions for report in reports])
        total_solar_savings = sum([report.co2_emissions_solar for report in reports])
        total_cost_savings = sum([report.cost_savings for report in reports])
        local_business_count = sum([1 for report in reports if report.type.title() == "Local Business"])
        council_facility_count = sum([1 for report in reports if report.type.title() == "Council Facility"])

        # For Council Facilities
        cf_total_ghg_savings = sum([report.total_emissions for report in reports if report.type.title() == "Council Facility"])
        cf_total_solar_savings = sum([report.co2_emissions_solar for report in reports if report.type.title() == "Council Facility"])
        cf_total_cost_savings = sum([report.cost_savings for report in reports if report.type.title() == "Council Facility"])

        # For Local Businesses
        lb_total_ghg_savings = sum([report.total_emissions for report in reports if report.type.title() == "Local Business"])
        lb_total_solar_savings = sum([report.co2_emissions_solar for report in reports if report.type.title() == "Local Business"])
        lb_total_cost_savings = sum([report.cost_savings for report in reports if report.type.title() == "Local Business"])

        return render_template('business_and_industry/view_reports.html', reporting_periods=reporting_periods, reports=reports, 
                                selected_period=reporting_period_id, total_ghg_savings=total_ghg_savings,
                                total_solar_savings=total_solar_savings, total_cost_savings=total_cost_savings, 
                                local_business_count=local_business_count, council_facility_count=council_facility_count,
                                cf_total_ghg_savings=cf_total_ghg_savings, cf_total_solar_savings=cf_total_solar_savings, 
                                cf_total_cost_savings=cf_total_cost_savings, lb_total_ghg_savings=lb_total_ghg_savings, 
                                lb_total_solar_savings=lb_total_solar_savings, lb_total_cost_savings=lb_total_cost_savings)

    @app.route('/update_report', methods=['POST'])
    def update_report():
        report_id = request.form.get('report_id')
        column_name = request.form.get('column_name')
        value = request.form.get('value')

        report = EmissionReport.query.get(report_id)

        if not report:
            return jsonify(status="error", message="Report not found")

        if column_name == "co2_emissions_solar":
            report.co2_emissions_solar = float(value)
        elif column_name == "cost_savings":
            report.cost_savings = float(value)
        elif column_name == "total_emissions":
            report.total_emissions = float(value)
        elif column_name == "type":
            report.business_or_facility.type = value
        else:
            return jsonify(status="error", message="Invalid column name")

        db.session.commit()

        return jsonify(status="success", other_emissions=report.other_emissions)

    @app.route('/delete_report/<int:report_id>', methods=['POST'])
    def delete_report(report_id):
        report = EmissionReport.query.get(report_id)
        
        if report:
            business_id = report.business_or_facility_id
            db.session.delete(report)
            db.session.commit()
            
            # Check if the business has other reports
            other_reports = EmissionReport.query.filter_by(business_or_facility_id=business_id).all()
            if not other_reports:  # If no other reports for this business
                business = BusinessOrFacility.query.get(business_id)
                if business:
                    db.session.delete(business)
                    db.session.commit()

            return jsonify(status="success")
        else:
            return jsonify(status="error", message="Report not found")

    @app.route('/add_business_report', methods=['GET', 'POST'])
    def add_business_report():
        businesses_or_facilities = BusinessOrFacility.query.all()  # Get all businesses
        reporting_periods = ReportingPeriod.query.all()
        show_error_modal = False
        error_msg = ""
        
        if request.method == 'POST':
            business_or_facility_id = request.form['business_or_facility']
            
            if business_or_facility_id == "new":  # if the user wants to add a new business
                name = request.form['name']
                type_ = request.form['type']

                existing_biz = BusinessOrFacility.query.filter_by(name=name).first()
                if existing_biz:
                    error_msg = "This business or facility already exists!"
                    show_error_modal = True
                    return render_template('business_and_industry/add_business_report.html', businesses_or_facilities=businesses_or_facilities, reporting_periods=reporting_periods, show_error_modal=show_error_modal, error_msg=error_msg)
                else:
                    business_or_facility = BusinessOrFacility(name=name, type=type_)
                    db.session.add(business_or_facility)
                    db.session.commit()
                    business_or_facility_id = business_or_facility.id  # Update the id for new business
            
            reporting_period_id = request.form['reporting_period']

            existing_report = EmissionReport.query.filter_by(business_or_facility_id=business_or_facility_id, reporting_period_id=reporting_period_id).first()
            if existing_report:
                error_msg = "An emission report for this business in the selected reporting period already exists!"
                show_error_modal = True
            else:
                co2_emissions_solar = request.form['co2_emissions_solar']
                total_emissions = request.form['total_emissions']
                cost_savings = request.form['cost_savings']

                report = EmissionReport(co2_emissions_solar=co2_emissions_solar, total_emissions=total_emissions, 
                                        cost_savings=cost_savings, reporting_period_id=reporting_period_id, 
                                        business_or_facility_id=business_or_facility_id)
                db.session.add(report)
                db.session.commit()
                return redirect(url_for('view_reports'))

        return render_template('business_and_industry/add_business_report.html', businesses_or_facilities=businesses_or_facilities, reporting_periods=reporting_periods, show_error_modal=show_error_modal, error_msg=error_msg)

    @app.route('/api/emissions_by_reporting_period')
    def api_emissions_by_reporting_period():
        periods = ReportingPeriod.query.all()
        data = []
        for period in periods:
            reports = period.reports
            total_co2_solar = sum([r.co2_emissions_solar for r in reports])
            total_other = sum([r.other_emissions for r in reports])
            data.append({
                'start_month': period.start_month,
                'end_month': period.end_month,
                'start_year': period.start_year,
                'end_year': period.end_year,
                'co2_solar': total_co2_solar,
                'other_emissions': total_other
            })
        return jsonify(data)
    
    @app.route('/api/emissions_by_council_and_business')
    def api_emissions_by_council_and_business():
        businesses_and_facilities = BusinessOrFacility.query.all()
        data = []

        for entity in businesses_and_facilities:
            reports = entity.reports
            total_co2_solar = sum([r.co2_emissions_solar for r in reports])
            total_other = sum([r.other_emissions for r in reports])

            data.append({
                'name': entity.name,
                'type': entity.type,  # This will be either 'local business' or 'council facility'
                'co2_solar': total_co2_solar,
                'other_emissions': total_other
            })

        return jsonify(data)

    @app.route('/api/top5_council_facilities_by_total_emissions')
    def top5_council_facilities_by_total_emissions():
        facilities = BusinessOrFacility.query.filter_by(type='council facility').all()
        data = []
        for facility in facilities:
            reports = facility.reports
            total_emissions = sum([r.total_emissions for r in reports])
            data.append({
                'name': facility.name,
                'co2_solar': sum([r.co2_emissions_solar for r in reports]),
                'other_emissions': sum([r.other_emissions for r in reports]),
                'total_emissions': total_emissions
            })
        data.sort(key=lambda x: x['total_emissions'], reverse=True)
        return jsonify(data[:5])

    @app.route('/api/top5_local_businesses_by_total_emissions')
    def top5_local_businesses_by_total_emissions():
        facilities = BusinessOrFacility.query.filter_by(type='local business').all()
        data = []
        for facility in facilities:
            reports = facility.reports
            total_emissions = sum([r.total_emissions for r in reports])
            data.append({
                'name': facility.name,
                'co2_solar': sum([r.co2_emissions_solar for r in reports]),
                'other_emissions': sum([r.other_emissions for r in reports]),
                'total_emissions': total_emissions
            })
        data.sort(key=lambda x: x['total_emissions'], reverse=True)
        return jsonify(data[:5])

    @app.route('/api/top5_council_facilities_by_solar_pv')
    def top5_council_facilities_by_solar_pv():
        facilities = BusinessOrFacility.query.filter_by(type='council facility').all()
        data = []
        for facility in facilities:
            reports = facility.reports
            total_solar = sum([r.co2_emissions_solar for r in reports])
            data.append({
                'name': facility.name,
                'co2_solar': total_solar,
                'other_emissions': sum([r.other_emissions for r in reports])
            })
        data.sort(key=lambda x: x['co2_solar'], reverse=True)
        return jsonify(data[:5])

    @app.route('/api/top5_local_businesses_by_solar_pv')
    def top5_local_businesses_by_solar_pv():
        facilities = BusinessOrFacility.query.filter_by(type='local business').all()
        data = []
        for facility in facilities:
            reports = facility.reports
            total_solar = sum([r.co2_emissions_solar for r in reports])
            data.append({
                'name': facility.name,
                'co2_solar': total_solar,
                'other_emissions': sum([r.other_emissions for r in reports])
            })
        data.sort(key=lambda x: x['co2_solar'], reverse=True)
        return jsonify(data[:5])

    @app.route('/api/top5_council_facilities_by_other_emissions')
    def top5_council_facilities_by_other_emissions():
        facilities = BusinessOrFacility.query.filter_by(type='council facility').all()
        data = []
        for facility in facilities:
            reports = facility.reports
            other_emissions = sum([r.other_emissions for r in reports])
            data.append({
                'name': facility.name,
                'co2_solar': sum([r.co2_emissions_solar for r in reports]),
                'other_emissions': other_emissions
            })
        data.sort(key=lambda x: x['other_emissions'], reverse=True)
        return jsonify(data[:5])

    @app.route('/api/top5_local_businesses_by_other_emissions')
    def top5_local_businesses_by_other_emissions():
        facilities = BusinessOrFacility.query.filter_by(type='local business').all()
        data = []
        for facility in facilities:
            reports = facility.reports
            other_emissions = sum([r.other_emissions for r in reports])
            data.append({
                'name': facility.name,
                'co2_solar': sum([r.co2_emissions_solar for r in reports]),
                'other_emissions': other_emissions
            })
        data.sort(key=lambda x: x['other_emissions'], reverse=True)
        return jsonify(data[:5])

    @app.route('/api/top5_facilities_by_cost_savings')
    def top5_facilities_by_cost_savings():
        facilities = BusinessOrFacility.query.all()
        data = []
        for facility in facilities:
            reports = facility.reports
            cost_savings = sum([r.cost_savings for r in reports])
            data.append({
                'name': facility.name,
                'cost_savings': cost_savings
            })
        data.sort(key=lambda x: x['cost_savings'], reverse=True)
        return jsonify(data[:5])

    @app.route('/api/bottom5_facilities_by_cost_savings')
    def bottom5_facilities_by_cost_savings():
        facilities = BusinessOrFacility.query.all()
        data = []
        for facility in facilities:
            reports = facility.reports
            cost_savings = sum([r.cost_savings for r in reports])
            data.append({
                'name': facility.name,
                'cost_savings': cost_savings
            })
        data.sort(key=lambda x: x['cost_savings'])
        return jsonify(data[:5])
    
    @app.route('/download_excel')
    def download_excel():
        wb = openpyxl.Workbook()
        
        # Define the headers
        headers = ["Name", "Type", "Start Month", "End Month", "Start Year", "End Year", 
                "CO2 Emissions (Solar)", "Total Emissions", "Cost Savings", "Other Emissions"]

        # Remove the default sheet created and start with a fresh workbook
        wb.remove(wb.active)

        # Loop through each unique reporting period
        for period in ReportingPeriod.query.all():
            # Create a new sheet for this reporting period
            ws = wb.create_sheet(title=f"{period.start_month} {period.start_year} - {period.end_month} {period.end_year}")

            # Write headers
            for col_num, header in enumerate(headers, 1):
                col_letter = get_column_letter(col_num)
                ws['{}1'.format(col_letter)] = header
                ws.column_dimensions[col_letter].width = len(header) + 5

            # Write the data for this reporting period
            row_num = 2
            for report in EmissionReport.query.filter_by(reporting_period_id=period.id).all():
                values = [
                    report.business_or_facility.name,
                    report.type,
                    report.reporting_period.start_month,
                    report.reporting_period.end_month,
                    report.reporting_period.start_year,
                    report.reporting_period.end_year,
                    report.co2_emissions_solar,
                    report.total_emissions,
                    report.cost_savings,
                    report.other_emissions
                ]
                for col_num, value in enumerate(values, 1):
                    ws.cell(row=row_num, column=col_num, value=value)
                row_num += 1

        # Save the workbook to a bytes buffer
        bytes_io = BytesIO()
        wb.save(bytes_io)
        bytes_io.seek(0)

        return send_file(bytes_io, download_name="emission_reports.xlsx", as_attachment=True)

    @app.route('/download_word', methods=['POST'])
    def download_word():
        data = request.json  # assuming you send the data as JSON
        images_data = data['images']

        # Create a new document
        doc = Document()
        
        for img_data in images_data:
            title = img_data['title']
            img_base64 = img_data['data']

            # Add the title
            doc.add_heading(title, level=1)

            # Convert base64 image to bytes and add it to the document
            img_stream = BytesIO(base64.b64decode(img_base64.split(',')[1]))
            doc.add_picture(img_stream, width=Inches(6))

        # Save the document to bytes buffer
        bytes_io = BytesIO()
        doc.save(bytes_io)
        bytes_io.seek(0)

        return send_file(bytes_io, download_name="graphs.docx", as_attachment=True)
    
    @app.route('/enter_energy_rating_data', methods=['GET', 'POST'])
    def enter_energy_rating_data():
        if request.method == 'POST':
            # process submitted form data
            class_ = request.form['class']
            year = request.form['year']
            half_year = request.form['half_year']
            star_rating = request.form['star_rating']
            certificates_issued = request.form['certificates_issued']
            avg_conditioned_area = request.form['avg_conditioned_area']

            # Get the emission factor for the given year
            emission_factor = EmissionFactor.query.filter_by(year=year).first()

            # create new EnergyRatingData object with emission_factor_id
            data = EnergyRatingData(
                class_=class_,
                year=year,
                half_year=half_year,
                star_rating=star_rating,
                certificates_issued=certificates_issued,
                avg_conditioned_area=avg_conditioned_area,
                emission_factor_id=emission_factor.id if emission_factor else None
            )

            db.session.add(data)
            db.session.commit()
            return redirect(url_for('view_energy_rating_data'))

        return render_template('esd/enter_energy_rating_data.html')
    
    @app.route('/enter_emission_factors_data', methods=['GET', 'POST'])
    def enter_emission_factors_data():
        if request.method == 'POST':
            year = request.form['year']
            factor = request.form['factor']

            # Check for existing emission factor for the year
            existing_factor = EmissionFactor.query.filter_by(year=year).first()

            if existing_factor:
                # Update the factor if it exists
                existing_factor.factor = factor
            else:
                # Create a new entry if it doesn't
                emission_factor = EmissionFactor(year=year, factor=factor)
                db.session.add(emission_factor)
            
            db.session.commit()
            return redirect(url_for('view_energy_rating_data'))

        return render_template('esd/enter_emission_factors_data.html')

    @app.route('/view_energy_rating_data', methods=['GET', 'POST'])
    def view_energy_rating_data():
        energy_data = EnergyRatingData.query.order_by(asc(EnergyRatingData.year), asc(EnergyRatingData.half_year)).all()

        if request.method == 'POST':
            # Implement your editing logic here, e.g., update the database records
            data_id = request.form.get('data_id')
            updated_value = request.form.get('updated_value')
            # Perform DB update operation here
            return jsonify(success=True)

        return render_template('esd/view_energy_rating_data.html', energy_data=energy_data)

    @app.route('/view_energy_rating_data/<int:data_id>', methods=['GET'])
    def view_energy_rating_data_with_id(data_id):
        energy_data = EnergyRatingData.query.order_by(
            asc(EnergyRatingData.year),
            asc(EnergyRatingData.half_year)
        ).all()

        return render_template(
            'esd/view_energy_rating_data.html',
            energy_data=energy_data,
            highlighted_id=data_id
        )

    @app.route('/delete_energy_data/<int:data_id>', methods=['POST'])
    def delete_energy_data(data_id):
        try:
            data = EnergyRatingData.query.get(data_id)
            if data:
                db.session.delete(data)
                db.session.commit()
                return jsonify(status='success')
            else:
                return jsonify(status='fail', message='Data not found')
        except Exception as e:
            return jsonify(status='fail', message=str(e))

    @app.route('/update_data', methods=['POST'])
    def update_data():
        data_id = request.form.get('data_id')
        column_name = request.form.get('column_name')
        value = request.form.get('value')

        if data_id is not None:
            data_id = int(data_id)
        else:
            print("Data ID is None")  # Print a message if data_id is None
            return jsonify(status="error", message="Data ID is None")

        data = EnergyRatingData.query.get(data_id)  # Removed the int conversion here

        if not data:
            return jsonify(status="error", message="Data not found")

        try:
            if column_name == "year":
                data.year = int(value)
            elif column_name == "half_year":
                data.half_year = int(value)
            elif column_name == "class":
                data.class_ = int(value)
            elif column_name == "star_rating":
                data.star_rating = float(value)
            elif column_name == "avg_conditioned_area":
                data.avg_conditioned_area = float(value)
            elif column_name == "certificates_issued":
                data.certificates_issued = int(value)
            else:
                return jsonify(status="error", message="Invalid column name")
        except ValueError:
            return jsonify(status="error", message="Invalid value format")

        db.session.commit()

        return jsonify(
            status="success",
            year=data.year,
            half_year=data.half_year,
            class_=data.class_,
            star_rating=data.star_rating,
            avg_conditioned_area=data.avg_conditioned_area,
            mj_saved_per_annum='{:,.2f}'.format(data.mj_saved_per_annum) if data.mj_saved_per_annum is not None else None,
            emissions_reduction='{:,.2f}'.format(data.emissions_reduction) if data.emissions_reduction is not None else None
        )

    @app.route('/api/mj_saved_per_annum_over_time', methods=['GET'])
    def get_mj_saved_per_annum_over_time():
        data = EnergyRatingData.query.all()

        # Nested defaultdict to support year, class, and half_year
        organized_data = defaultdict(lambda: defaultdict(lambda: defaultdict(list)))

        for item in data:
            year = item.year
            class_ = item.class_
            emissions = item.emissions_reduction
            mj_saved = item.mj_saved_per_annum
            half_year = item.half_year
            star_rating = item.star_rating
            id_ = item.id

            if mj_saved is not None:
                emissions_to_add = emissions if emissions is not None else 0
                organized_data[year][class_][half_year].append((mj_saved, id_, star_rating, emissions_to_add))


        # Calculating the average MJ saved per annum for each class per year
        results = []

        for year, classes in organized_data.items():
            for class_, half_years in classes.items():
                for half_year, mj_saved_list in half_years.items():
                    # Add the emissions as well as the mj
                    if mj_saved_list:   
                        avg_mj_saved = sum([item[0] for item in mj_saved_list]) / len(mj_saved_list)
                        avg_emissions = sum([item[3] for item in mj_saved_list]) / len(mj_saved_list)
                        id_ = mj_saved_list[0][1]  # Here, extract the id from the first tuple in the list (adjust as needed)
                        star_rating = mj_saved_list[0][2]

                        results.append({
                            "year": year,
                            "class": class_,
                            "avg_mj_saved_per_annum": avg_mj_saved,
                            "avg_emissions_reduction": avg_emissions,
                            "half_year": half_year,
                            "id": id_,  # Here, added the id to the result dictionary
                            "star_rating": star_rating
                        })


        return jsonify(results)

    @app.route('/upload_google_insights', methods=['POST'])
    def upload_google_insights():
        if 'file' not in request.files:
            flash('No file part')
            return redirect(request.url)

        file = request.files['file']

        if file.filename == '':
            flash('No selected file')
            return redirect(request.url)

        if file and allowed_file(file.filename):
            # Always save to the same file name to ensure only one file is saved
            filename = os.path.join(app.config['UPLOAD_FOLDER'], 'google_insights.csv')
            file.save(filename)

            with open('original_google_insights_name.txt', 'w') as f:
                f.write(file.filename)

            return redirect(url_for('ev_index'))

        return redirect(url_for('ev_index'))

    @app.route('/clear_google_insights', methods=['POST'])
    def clear_google_insights():
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], 'google_insights.csv')
        
        if os.path.exists(filepath):
            os.remove(filepath)

        return redirect(url_for('ev_index'))

    @app.route('/vista_summary', methods=['GET'])
    def vista_summary():
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], 'vista_data.csv')

        if os.path.exists(file_path):
            summary = process_vista_data()
            print(summary)
            return render_template('ev/vista_summary.html', vista_summary=summary)
        else:
            flash('No data uploaded yet. Please upload the VISTA data CSV file.')
            return redirect(url_for('ev_index'))

    @app.route('/ev_index')
    def ev_index():
        google_insights_filename = get_original_filename('original_google_insights_name.txt')
        vista_filename = get_original_filename('original_vista_name.txt')
        return render_template('ev/ev_index.html', google_insights_filename=google_insights_filename, vista_filename=vista_filename)
    
    def allowed_file(filename):
        return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

    @app.route('/upload_vista', methods=['POST'])
    def upload_vista():
        if 'file' not in request.files:
            flash('No file part')
            return redirect(request.url)

        file = request.files['file']

        if file.filename == '':
            flash('No selected file')
            return redirect(request.url)

        if file and allowed_file(file.filename):
            # Always save to the same file name to ensure only one file is saved
            filename = os.path.join(app.config['UPLOAD_FOLDER'], 'vista_data.csv')
            file.save(filename)

            with open('original_vista_name.txt', 'w') as f:
                f.write(file.filename)

            return redirect(url_for('ev_index'))

        return redirect(url_for('ev_index'))

    @app.route('/clear_vista', methods=['POST'])
    def clear_vista():
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], 'vista_data.csv')
        
        if os.path.exists(filepath):
            os.remove(filepath)

        return redirect(url_for('ev_index'))

    @app.route('/filtered_data', methods=['GET'])
    def filtered_data():
        year_filter = request.args.get('year_filter', 'all')
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], 'google_insights.csv')
        
        if os.path.exists(file_path):
            # Load the entire dataset to get unique years
            df_all = pd.read_csv(file_path)
            unique_years_all = df_all['year'].unique().tolist()

            # Load the filtered dataset for the summary
            summary = process_google_insights_data(file_path, year_filter)

            # Overwrite unique_years in summary with the list computed from entire dataset
            summary['unique_years'] = unique_years_all
            
            return render_template('ev/google_insights_summary.html', google_insights_summary=summary, year_filter=year_filter)
        else:
            flash('No data uploaded yet. Please upload the Google Insights Explorer data CSV file.')
            return redirect(url_for('ev_index'))
        
    def process_google_insights_data(filepath, year_filter='all'):
        try:
            df = pd.read_csv(filepath)

            if year_filter != 'all':
                df = df[df['year'] == int(year_filter)]
            else:
                year_filter = 'all'

            # Ensuring existence of required columns
            required_columns = ['year', 'mode', 'travel_bounds', 'trips', 'gpc_distance_km', 'gpc_co2e_tons']
            if not all(column in df.columns for column in required_columns):
                return {}
            df = df[df['travel_bounds'] != 'TOTAL']

            # CO2 emissions each year, broken down by mode of transport
            emissions_by_year_mode = df.groupby(['year', 'mode'])['gpc_co2e_tons'].sum().unstack().fillna(0)

            # Average emissions per trip for each mode of transport
            avg_emissions_per_trip = (df.groupby('mode')['gpc_co2e_tons'].sum() / df.groupby('mode')['trips'].sum()).fillna(0)

            # Emissions data by travel bounds
            emissions_by_bounds = df.groupby('travel_bounds')['gpc_co2e_tons'].sum()

            # CO2 emissions per kilometer for each mode of transport
            emissions_per_km = (df.groupby('mode')['gpc_co2e_tons'].sum() / df.groupby('mode')['gpc_distance_km'].sum()).fillna(0)

            summary = {
                'emissions_by_year_mode': emissions_by_year_mode.to_dict(),
                'avg_emissions_per_trip': avg_emissions_per_trip.to_dict(),
                'emissions_by_bounds': emissions_by_bounds.to_dict(),
                'emissions_per_km': emissions_per_km.to_dict(),
                'zero_emission_modes': df[df['gpc_co2e_tons'] == 0]['mode'].unique().tolist()  # Modes like cycling, walking contributing zero emissions
            }

            first_year_modes = list(next(iter(summary['emissions_by_year_mode'].values())).keys())

            summary.update({
                'first_year_modes': first_year_modes
            })

                    # Yearly emissions
            yearly_emissions = df.groupby('year')['gpc_co2e_tons'].sum().to_dict()

            # Percent of total transportation emissions
            total_emissions = df['gpc_co2e_tons'].sum()
            emissions_percent_by_bounds = (df.groupby('travel_bounds')['gpc_co2e_tons'].sum() / total_emissions * 100).to_dict()

            # Percent of total kilometers traveled
            total_km = df['gpc_distance_km'].sum()
            km_percent_by_bounds = (df.groupby('travel_bounds')['gpc_distance_km'].sum() / total_km * 100).to_dict()

            # Total combined number of trips per year
            total_trips_by_year = df.groupby('year')['trips'].sum().to_dict()

            # Total combined vehicle kilometers traveled per year
            total_vehicle_km_by_year = df.groupby('year')['gpc_distance_km'].sum().to_dict()

            total_vehicle_km_by_mode = df.groupby('mode')['gpc_distance_km'].sum().to_dict()


            # Percent of total combined kilometers by mode
            km_percent_by_mode = (df.groupby('mode')['gpc_distance_km'].sum() / total_km * 100).to_dict()
            emissions_percent_by_mode = (df.groupby('mode')['gpc_co2e_tons'].sum() / total_emissions * 100).to_dict()

            # Emissions by travel bounds and year
            emissions_by_bounds_year = df.groupby(['travel_bounds', 'year'])['gpc_co2e_tons'].sum().unstack().fillna(0).to_dict()

            unique_years = df['year'].unique().tolist()

            # Total trips by mode and year
            total_trips_by_mode_year = df.groupby(['mode', 'year'])['trips'].sum().unstack().fillna(0).to_dict()

            # Total kilometers by mode and year
            total_km_by_mode_year = df.groupby(['mode', 'year'])['gpc_distance_km'].sum().unstack().fillna(0).to_dict()

            unique_modes = df['mode'].unique().tolist()

            # Adding new data to summary dictionary
            summary.update({
                'yearly_emissions': yearly_emissions,
                'emissions_percent_by_bounds': emissions_percent_by_bounds,
                'km_percent_by_bounds': km_percent_by_bounds,
                'total_trips_by_year': total_trips_by_year,
                'km_percent_by_mode': km_percent_by_mode,
                'total_vehicle_km_by_year': total_vehicle_km_by_year,
                'emissions_by_bounds_year': emissions_by_bounds_year,
                'unique_years': unique_years,
                'total_trips_by_mode_year': total_trips_by_mode_year,
                'total_km_by_mode_year': total_km_by_mode_year,
                'total_vehicle_km_by_mode': total_vehicle_km_by_mode,
                'unique_modes': unique_modes,
                'emissions_percent_by_mode': emissions_percent_by_mode,
                'year_filter': year_filter
            })

            return summary
        except Exception as e:
            print("Error inside process_google_insights_data:", e)
            return {}

    @app.route('/google_insights_summary', methods=['GET'])
    def google_insights_summary():
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], 'google_insights.csv')

        if os.path.exists(file_path):
            summary = process_google_insights_data(file_path)
            return render_template('ev/google_insights_summary.html', google_insights_summary=summary)
        else:
            flash('No data uploaded yet. Please upload the Google Insights Explorer data CSV file.')
            return redirect(url_for('ev_index'))
        
    @app.route('/upload_evie', methods=['POST'])
    def upload_evie():
        if 'file' not in request.files:
            flash('No file part')
            return redirect(request.referrer)
        
        file = request.files['file']

        if file.filename == '':
            flash('No selected file')
            return redirect(request.referrer)
        
        if file:
            df = pd.read_csv(file)
            # TODO: Add your logic here to process the data frame as needed
            flash('Evie data uploaded successfully!')
            return redirect(url_for('ev_index'))

    return app

if __name__ == '__main__':
    app = create_app()
    app.run(debug=True)
    with app.app_context():
        db.create_all()
    